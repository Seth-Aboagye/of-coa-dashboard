import io
import re
from datetime import datetime

import numpy as np
import pandas as pd
import streamlit as st
import plotly.express as px

from docx import Document
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, landscape
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak


st.set_page_config(
    page_title="Oracle Fusion GL / COA Audit Dashboard",
    layout="wide",
    initial_sidebar_state="expanded"
)


# ============================================================
# Password Protection
# ============================================================

def check_password():
    if "authenticated" not in st.session_state:
        st.session_state["authenticated"] = False

    if st.session_state["authenticated"]:
        return True

    st.title("Secure Oracle Fusion COA Dashboard")
    password = st.text_input("Enter dashboard password", type="password")

    app_password = st.secrets.get("APP_PASSWORD", "")

    if password:
        if password == app_password:
            st.session_state["authenticated"] = True
            st.rerun()
        else:
            st.error("Incorrect password.")
            return False

    return False


if not check_password():
    st.stop()


# ============================================================
# App Header
# ============================================================

st.title("Oracle Fusion GL / COA Audit Analytics Dashboard")

st.caption(
    "Upload Oracle Fusion transaction exports and the COA workbook to decode, analyze, "
    "flag exceptions, and export reports."
)

st.info(
    "Privacy note: Uploaded files are processed during the active session only. "
    "This app does not intentionally save Oracle exports to a database or permanent storage."
)


# ============================================================
# Helper Functions
# ============================================================

def normalize_col(col):
    return str(col).strip().lower().replace("\n", " ").replace("_", " ")


def find_col(df, possible_names):
    normalized = {normalize_col(c): c for c in df.columns}
    possible_names = [normalize_col(x) for x in possible_names]

    for name in possible_names:
        if name in normalized:
            return normalized[name]

    for norm, original in normalized.items():
        for name in possible_names:
            if name in norm:
                return original

    return None


def clean_code(value, width=None):
    if pd.isna(value):
        return ""

    value = str(value).strip()

    if value.endswith(".0"):
        value = value[:-2]

    value = re.sub(r"[^A-Za-z0-9]", "", value)

    if width and value.isdigit():
        value = value.zfill(width)

    return value


def read_excel_or_csv(uploaded_file):
    name = uploaded_file.name.lower()

    if name.endswith(".csv"):
        return pd.read_csv(uploaded_file)

    return pd.read_excel(uploaded_file)


def extract_lookup_from_sheet(df, sheet_name):
    if df is None or df.empty:
        return pd.DataFrame(columns=["Code", "Description", "Source Tab"])

    df = df.copy()
    df.columns = [str(c).strip() for c in df.columns]

    description_col = find_col(df, ["Description", "Name", "Account Description"])

    if description_col is None:
        return pd.DataFrame(columns=["Code", "Description", "Source Tab"])

    level_cols = [c for c in df.columns if "level" in normalize_col(c)]

    records = []

    for _, row in df.iterrows():
        desc = row.get(description_col)

        if pd.isna(desc) or str(desc).strip().lower() in ["nan", ""]:
            continue

        for col in level_cols:
            code = clean_code(row.get(col))

            if code == "":
                continue

            if "Z" in code.upper():
                continue

            records.append(
                {
                    "Code": code,
                    "Description": str(desc).strip(),
                    "Source Tab": sheet_name,
                }
            )

    if not records:
        return pd.DataFrame(columns=["Code", "Description", "Source Tab"])

    return pd.DataFrame(records).drop_duplicates()


def load_coa(coa_file):
    xls = pd.ExcelFile(coa_file)

    expected_tabs = [
        "Fund",
        "Business Function",
        "Cost Center",
        "Agency",
        "Account",
        "Program",
        "Inter Fund",
        "Project",
        "Budget Year",
        "Future",
    ]

    lookups = {}

    for tab in expected_tabs:
        if tab in xls.sheet_names:
            raw = pd.read_excel(coa_file, sheet_name=tab)
            lookups[tab] = extract_lookup_from_sheet(raw, tab)
        else:
            lookups[tab] = pd.DataFrame(columns=["Code", "Description", "Source Tab"])

    return lookups


def detect_account_string_col(df):
    candidates = [
        "account string",
        "account combination",
        "code combination",
        "accounting flexfield",
        "concatenated segments",
        "gl account",
        "account code",
        "coa",
        "combination",
    ]

    return find_col(df, candidates)


def split_account_string(df, account_col):
    df = df.copy()

    if account_col is None:
        return df

    parts = (
        df[account_col]
        .astype(str)
        .str.replace(".", "-", regex=False)
        .str.split("-", expand=True)
    )

    segment_names = [
        "Fund",
        "Business Function",
        "Cost Center",
        "Account",
        "Program",
        "Agency",
        "Inter Fund",
        "Project",
        "Budget Year",
        "Future",
    ]

    widths = {
        "Fund": 4,
        "Business Function": 2,
        "Cost Center": 6,
        "Account": 6,
        "Program": 4,
        "Agency": 4,
        "Inter Fund": 4,
        "Project": 6,
        "Budget Year": 2,
        "Future": 4,
    }

    for i, segment in enumerate(segment_names):
        if i < parts.shape[1]:
            df[segment] = parts[i].apply(lambda x: clean_code(x, widths.get(segment)))
        else:
            if segment not in df.columns:
                df[segment] = ""

    return df


def ensure_segment_columns(df):
    df = df.copy()

    mapping = {
        "Fund": ["fund", "fund segment", "segment1"],
        "Business Function": ["business function", "function", "bf", "segment2"],
        "Cost Center": ["cost center", "department", "dept", "segment3"],
        "Account": ["account", "natural account", "account segment", "segment4"],
        "Program": ["program", "program segment", "segment5"],
        "Agency": ["agency", "segment6"],
        "Inter Fund": ["inter fund", "interfund", "segment7"],
        "Project": ["project", "project number", "segment8"],
        "Budget Year": ["budget year", "fiscal year", "segment9"],
        "Future": ["future", "segment10"],
    }

    widths = {
        "Fund": 4,
        "Business Function": 2,
        "Cost Center": 6,
        "Account": 6,
        "Program": 4,
        "Agency": 4,
        "Inter Fund": 4,
        "Project": 6,
        "Budget Year": 2,
        "Future": 4,
    }

    for segment, names in mapping.items():
        if segment not in df.columns:
            col = find_col(df, names)
            if col is not None:
                df[segment] = df[col].apply(lambda x: clean_code(x, widths.get(segment)))
        else:
            df[segment] = df[segment].apply(lambda x: clean_code(x, widths.get(segment)))

    return df


def add_coa_descriptions(df, lookups):
    df = df.copy()

    for segment, lookup in lookups.items():
        if segment not in df.columns:
            continue

        desc_col = f"{segment} Description"

        if lookup.empty:
            df[desc_col] = ""
            continue

        lu = lookup[["Code", "Description"]].drop_duplicates("Code")

        df = df.merge(
            lu.rename(columns={"Code": segment, "Description": desc_col}),
            on=segment,
            how="left",
        )

    return df


def detect_amount_columns(df):
    df = df.copy()

    debit_col = find_col(df, ["debit", "entered debit", "accounted debit", "dr"])
    credit_col = find_col(df, ["credit", "entered credit", "accounted credit", "cr"])
    amount_col = find_col(df, ["amount", "net amount", "transaction amount", "accounted amount"])

    if debit_col:
        df[debit_col] = pd.to_numeric(df[debit_col], errors="coerce").fillna(0)

    if credit_col:
        df[credit_col] = pd.to_numeric(df[credit_col], errors="coerce").fillna(0)

    if debit_col and credit_col:
        df["Debit"] = df[debit_col]
        df["Credit"] = df[credit_col]
        df["Net Amount"] = df["Debit"] - df["Credit"]
    elif amount_col:
        df[amount_col] = pd.to_numeric(df[amount_col], errors="coerce").fillna(0)
        df["Net Amount"] = df[amount_col]
        df["Debit"] = np.where(df["Net Amount"] >= 0, df["Net Amount"], 0)
        df["Credit"] = np.where(df["Net Amount"] < 0, -df["Net Amount"], 0)
    else:
        df["Debit"] = 0
        df["Credit"] = 0
        df["Net Amount"] = 0

    df["Absolute Amount"] = df["Net Amount"].abs()

    return df


def detect_date_col(df):
    df = df.copy()

    date_col = find_col(
        df,
        [
            "accounting date",
            "effective date",
            "journal date",
            "transaction date",
            "date",
            "posted date",
            "gl date",
        ],
    )

    if date_col:
        df["Transaction Date"] = pd.to_datetime(df[date_col], errors="coerce")
    else:
        df["Transaction Date"] = pd.NaT

    return df


def detect_description_col(df):
    return find_col(
        df,
        [
            "description",
            "journal description",
            "line description",
            "transaction description",
            "memo",
        ],
    )


def append_flag(existing, new_flag):
    existing = "" if pd.isna(existing) else str(existing).strip()
    if existing == "":
        return new_flag
    if new_flag in existing:
        return existing
    return existing + "; " + new_flag


def add_exceptions(
    df,
    large_threshold,
    round_amount_threshold=10000,
    after_hours_start=20,
    after_hours_end=6,
    dormant_days=180,
):
    df = df.copy()

    desc_col = detect_description_col(df)

    inputter_col = find_col(
        df,
        [
            "inputter",
            "entered by",
            "created by",
            "prepared by",
            "journal creator",
            "user name",
            "created user",
            "creator",
        ],
    )

    approver_col = find_col(
        df,
        [
            "approver",
            "approved by",
            "authorized by",
            "authorizer",
            "journal approver",
            "approval user",
            "approved user",
        ],
    )

    batch_col = find_col(
        df,
        [
            "batch id",
            "batch number",
            "batch name",
            "journal batch",
            "batch",
            "journal batch name",
        ],
    )

    source_col = find_col(df, ["journal source", "source", "transaction source"])
    category_col = find_col(df, ["journal category", "category"])

    invoice_col = find_col(
        df,
        [
            "invoice number",
            "invoice num",
            "invoice",
            "voucher number",
            "voucher",
            "supplier invoice",
        ],
    )

    budget_col = find_col(
        df,
        ["budget", "budget amount", "approved budget", "original budget", "current budget"],
    )

    actual_col = find_col(
        df,
        ["actual", "actual amount", "expenditure", "expense amount", "ytd actual"],
    )

    available_budget_col = find_col(
        df,
        ["available budget", "remaining budget", "budget remaining", "funds available"],
    )

    segment_cols = ["Fund", "Business Function", "Cost Center", "Account", "Program"]

    flags = []

    for _, row in df.iterrows():
        row_flags = []

        # Missing COA segment
        for col in segment_cols:
            if col in df.columns and str(row.get(col, "")).strip() == "":
                row_flags.append(f"Missing {col}")

        # Unknown COA code
        for col in segment_cols:
            desc_col_segment = f"{col} Description"
            if desc_col_segment in df.columns:
                if str(row.get(col, "")).strip() != "" and pd.isna(row.get(desc_col_segment)):
                    row_flags.append(f"Unknown {col} code")

        # Missing description
        if desc_col and str(row.get(desc_col, "")).strip() in ["", "nan", "None"]:
            row_flags.append("Missing transaction description")

        # Zero amount
        if abs(row.get("Net Amount", 0)) == 0:
            row_flags.append("Zero amount transaction")

        # Large transaction
        if abs(row.get("Net Amount", 0)) >= large_threshold:
            row_flags.append("Large transaction")

        tx_date = row.get("Transaction Date")

        # Future-dated transaction
        if pd.notna(tx_date) and tx_date.date() > datetime.today().date():
            row_flags.append("Future-dated transaction")

        # Weekend posting
        if pd.notna(tx_date) and tx_date.weekday() >= 5:
            row_flags.append("Weekend posting")

        # After-hours posting
        if pd.notna(tx_date):
            hour = tx_date.hour
            if hour >= after_hours_start or hour < after_hours_end:
                row_flags.append("After-hours posting")

        # Manual journal entry
        source_value = str(row.get(source_col, "")).lower() if source_col else ""
        category_value = str(row.get(category_col, "")).lower() if category_col else ""

        if "manual" in source_value or "manual" in category_value:
            row_flags.append("Manual journal entry")

        # Suspicious round-dollar amount
        amount = abs(row.get("Net Amount", 0))
        if amount >= round_amount_threshold and amount % 1000 == 0:
            row_flags.append("Suspicious round-dollar amount")

        # Budget overrun
        if available_budget_col:
            available_budget = pd.to_numeric(row.get(available_budget_col), errors="coerce")
            if pd.notna(available_budget) and available_budget < 0:
                row_flags.append("Budget overrun")

        elif budget_col and actual_col:
            budget = pd.to_numeric(row.get(budget_col), errors="coerce")
            actual = pd.to_numeric(row.get(actual_col), errors="coerce")
            if pd.notna(budget) and pd.notna(actual) and actual > budget:
                row_flags.append("Budget overrun")

        # Segregation of duties issue
        if inputter_col and approver_col:
            inputter = str(row.get(inputter_col, "")).strip().lower()
            approver = str(row.get(approver_col, "")).strip().lower()

            if inputter and approver and inputter not in ["nan", "none"] and inputter == approver:
                row_flags.append("Segregation of duties issue: inputter also approved")

        flags.append("; ".join(row_flags))

    df["Exception Flags"] = flags

    # Duplicate invoice number
    if invoice_col and "Net Amount" in df.columns:
        duplicate_invoice_mask = df.duplicated(subset=[invoice_col, "Net Amount"], keep=False)
        df.loc[duplicate_invoice_mask, "Exception Flags"] = df.loc[
            duplicate_invoice_mask, "Exception Flags"
        ].apply(lambda x: append_flag(x, "Duplicate invoice number"))

    # Dormant account activity
    if "Account" in df.columns and df["Transaction Date"].notna().any():
        sorted_df = df.sort_values(["Account", "Transaction Date"]).copy()
        sorted_df["Previous Account Date"] = sorted_df.groupby("Account")["Transaction Date"].shift(1)
        sorted_df["Days Since Previous Activity"] = (
            sorted_df["Transaction Date"] - sorted_df["Previous Account Date"]
        ).dt.days

        dormant_indices = sorted_df.index[sorted_df["Days Since Previous Activity"] >= dormant_days]

        df.loc[dormant_indices, "Exception Flags"] = df.loc[dormant_indices, "Exception Flags"].apply(
            lambda x: append_flag(x, "Dormant account activity")
        )

    # Non-sequential batch IDs
    if batch_col:
        temp = df[[batch_col]].copy()
        temp["Batch Numeric"] = temp[batch_col].astype(str).str.extract(r"(\d+)$")
        temp["Batch Numeric"] = pd.to_numeric(temp["Batch Numeric"], errors="coerce")

        batch_numbers = sorted(temp["Batch Numeric"].dropna().unique())

        missing_batches = []
        if len(batch_numbers) > 1:
            expected = set(range(int(min(batch_numbers)), int(max(batch_numbers)) + 1))
            actual = set(int(x) for x in batch_numbers)
            missing_batches = sorted(expected - actual)

        if missing_batches:
            df.loc[df[batch_col].notna(), "Exception Flags"] = df.loc[
                df[batch_col].notna(), "Exception Flags"
            ].apply(lambda x: append_flag(x, "Batch ID sequence gap detected"))

            df["Missing Batch IDs"] = ", ".join(str(x) for x in missing_batches[:50])
        else:
            df["Missing Batch IDs"] = ""

    df["Has Exception"] = df["Exception Flags"].apply(lambda x: bool(str(x).strip()))

    return df


def add_duplicate_flags(df):
    df = df.copy()

    desc_col = detect_description_col(df)

    subset = [
        "Fund",
        "Cost Center",
        "Account",
        "Program",
        "Net Amount",
        "Transaction Date",
    ]

    subset = [c for c in subset if c in df.columns]

    if desc_col:
        subset.append(desc_col)

    if subset:
        df["Potential Duplicate"] = df.duplicated(subset=subset, keep=False)
    else:
        df["Potential Duplicate"] = False

    df.loc[df["Potential Duplicate"], "Exception Flags"] = df.loc[
        df["Potential Duplicate"], "Exception Flags"
    ].apply(lambda x: append_flag(x, "Potential duplicate"))

    df["Has Exception"] = df["Exception Flags"].apply(lambda x: bool(str(x).strip()))

    return df


def create_summary(df):
    rows = len(df)
    exception_count = int(df["Has Exception"].sum())
    duplicate_count = int(df["Potential Duplicate"].sum()) if "Potential Duplicate" in df.columns else 0

    return {
        "Total Rows": rows,
        "Total Debits": df["Debit"].sum(),
        "Total Credits": df["Credit"].sum(),
        "Net Amount": df["Net Amount"].sum(),
        "Gross Activity": df["Absolute Amount"].sum(),
        "Exception Count": exception_count,
        "Potential Duplicate Count": duplicate_count,
        "Exception Rate": exception_count / rows if rows else 0,
    }


def generate_conclusion(summary, top_dept, top_account):
    text = []

    text.append(
        f"The uploaded Oracle Fusion export contains {summary['Total Rows']:,} transaction lines "
        f"with gross activity of ${summary['Gross Activity']:,.2f}."
    )

    text.append(
        f"Total debits are ${summary['Total Debits']:,.2f}, total credits are "
        f"${summary['Total Credits']:,.2f}, and the net amount is "
        f"${summary['Net Amount']:,.2f}."
    )

    text.append(
        f"The dashboard identified {summary['Exception Count']:,} exception lines, "
        f"representing {summary['Exception Rate']:.1%} of total lines."
    )

    if summary["Potential Duplicate Count"] > 0:
        text.append(
            f"{summary['Potential Duplicate Count']:,} lines were flagged as potential duplicates."
        )

    if top_dept:
        text.append(f"The highest activity cost center is {top_dept}.")

    if top_account:
        text.append(f"The highest activity account is {top_account}.")

    text.append(
        "Recommended next steps: review large transactions, unknown COA codes, duplicate flags, "
        "future-dated postings, weekend postings, after-hours postings, manual journals, budget overruns, "
        "round-dollar transactions, dormant account activity, segregation of duties issues, duplicate invoices, "
        "and batch sequence gaps."
    )

    return "\n\n".join(text)


def export_excel(df, summary, conclusion, exception_summary):
    output = io.BytesIO()

    with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
        df.to_excel(writer, sheet_name="Decoded Transactions", index=False)
        df[df["Has Exception"]].to_excel(writer, sheet_name="Exceptions", index=False)
        pd.DataFrame([summary]).to_excel(writer, sheet_name="Summary KPIs", index=False)
        exception_summary.to_excel(writer, sheet_name="Exception Summary", index=False)
        pd.DataFrame({"Conclusion": [conclusion]}).to_excel(
            writer, sheet_name="Conclusion", index=False
        )

        workbook = writer.book
        header_fmt = workbook.add_format({"bold": True, "bg_color": "#D9EAF7"})
        money_fmt = workbook.add_format({"num_format": "$#,##0.00"})
        pct_fmt = workbook.add_format({"num_format": "0.0%"})

        for sheet_name, worksheet in writer.sheets.items():
            worksheet.set_row(0, None, header_fmt)
            worksheet.set_column(0, 40, 18)

        writer.sheets["Summary KPIs"].set_column(1, 5, 18, money_fmt)
        writer.sheets["Summary KPIs"].set_column(7, 7, 12, pct_fmt)

    output.seek(0)
    return output


def export_word(summary, conclusion, exceptions, exception_summary):
    output = io.BytesIO()

    doc = Document()
    doc.add_heading("Oracle Fusion GL / COA Audit Analytics Report", 0)
    doc.add_paragraph(f"Generated on {datetime.now().strftime('%B %d, %Y %I:%M %p')}")

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(conclusion)

    doc.add_heading("Key Performance Indicators", level=1)

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "KPI"
    table.rows[0].cells[1].text = "Value"

    for k, v in summary.items():
        row = table.add_row().cells
        row[0].text = str(k)

        if isinstance(v, float):
            row[1].text = f"{v:.1%}" if "Rate" in k else f"{v:,.2f}"
        else:
            row[1].text = f"{v:,}" if isinstance(v, int) else str(v)

    doc.add_heading("Exception Summary", level=1)

    if exception_summary.empty:
        doc.add_paragraph("No exceptions were identified.")
    else:
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        table.rows[0].cells[0].text = "Exception Type"
        table.rows[0].cells[1].text = "Count"

        for _, r in exception_summary.iterrows():
            row = table.add_row().cells
            row[0].text = str(r["Exception Type"])
            row[1].text = str(r["Count"])

    doc.add_heading("Exception Preview", level=1)

    if exceptions.empty:
        doc.add_paragraph("No exceptions were identified.")
    else:
        preview = exceptions.head(20)
        cols = list(preview.columns[:6])

        table = doc.add_table(rows=1, cols=len(cols))
        table.style = "Table Grid"

        for i, col in enumerate(cols):
            table.rows[0].cells[i].text = str(col)

        for _, r in preview[cols].iterrows():
            row = table.add_row().cells
            for i, col in enumerate(cols):
                row[i].text = str(r[col])

    doc.save(output)
    output.seek(0)

    return output


def export_pdf(summary, conclusion, exceptions, exception_summary):
    output = io.BytesIO()

    doc = SimpleDocTemplate(
        output,
        pagesize=landscape(letter),
        rightMargin=30,
        leftMargin=30,
        topMargin=30,
        bottomMargin=30,
    )

    styles = getSampleStyleSheet()
    story = []

    story.append(Paragraph("Oracle Fusion GL / COA Audit Analytics Report", styles["Title"]))
    story.append(
        Paragraph(
            f"Generated on {datetime.now().strftime('%B %d, %Y %I:%M %p')}",
            styles["Normal"],
        )
    )

    story.append(Spacer(1, 12))
    story.append(Paragraph("Executive Summary", styles["Heading1"]))

    for paragraph in conclusion.split("\n\n"):
        story.append(Paragraph(paragraph, styles["Normal"]))
        story.append(Spacer(1, 8))

    story.append(Paragraph("Key Performance Indicators", styles["Heading1"]))

    data = [["KPI", "Value"]]

    for k, v in summary.items():
        if isinstance(v, float):
            val = f"{v:.1%}" if "Rate" in k else f"{v:,.2f}"
        elif isinstance(v, int):
            val = f"{v:,}"
        else:
            val = str(v)

        data.append([k, val])

    table = Table(data, colWidths=[220, 180])
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.lightblue),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ]
        )
    )

    story.append(table)
    story.append(Spacer(1, 14))

    story.append(Paragraph("Exception Summary", styles["Heading1"]))

    if exception_summary.empty:
        story.append(Paragraph("No exceptions were identified.", styles["Normal"]))
    else:
        exception_data = [["Exception Type", "Count"]] + exception_summary.astype(str).values.tolist()
        exception_table = Table(exception_data, colWidths=[300, 100])
        exception_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.lightblue),
                    ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ]
            )
        )
        story.append(exception_table)

    story.append(PageBreak())
    story.append(Paragraph("Exception Preview", styles["Heading1"]))

    if exceptions.empty:
        story.append(Paragraph("No exceptions were identified.", styles["Normal"]))
    else:
        preview = exceptions.head(20)
        cols = [
            c
            for c in [
                "Transaction Date",
                "Fund",
                "Cost Center",
                "Account",
                "Program",
                "Net Amount",
                "Exception Flags",
            ]
            if c in preview.columns
        ]

        table_data = [cols] + preview[cols].astype(str).values.tolist()

        table = Table(table_data, repeatRows=1)
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.lightblue),
                    ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, -1), 7),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ]
            )
        )

        story.append(table)

    doc.build(story)
    output.seek(0)

    return output


def build_exception_summary(df):
    if "Exception Flags" not in df.columns:
        return pd.DataFrame(columns=["Exception Type", "Count"])

    flags = []

    for value in df["Exception Flags"].dropna():
        for flag in str(value).split(";"):
            flag = flag.strip()
            if flag:
                flags.append(flag)

    if not flags:
        return pd.DataFrame(columns=["Exception Type", "Count"])

    return (
        pd.Series(flags)
        .value_counts()
        .reset_index()
        .rename(columns={"index": "Exception Type", 0: "Count"})
    )


# ============================================================
# Sidebar
# ============================================================

st.sidebar.header("1. Upload Files")

oracle_file = st.sidebar.file_uploader(
    "Upload Oracle Fusion export",
    type=["xlsx", "xls", "csv"],
    help="Upload Oracle Fusion GL, AP, AR, payroll, budget, or journal export.",
)

coa_file = st.sidebar.file_uploader(
    "Upload COA workbook",
    type=["xlsx", "xls"],
    help="Upload the Chart of Accounts workbook.",
)

if st.sidebar.button("Clear uploaded files / reset session"):
    for key in list(st.session_state.keys()):
        del st.session_state[key]
    st.rerun()

st.sidebar.header("2. Exception Settings")

large_threshold = st.sidebar.number_input(
    "Large transaction threshold",
    min_value=0,
    value=100000,
    step=10000,
)

show_only_exceptions = st.sidebar.checkbox("Show only exceptions", value=False)

st.sidebar.header("3. Advanced Audit Settings")

round_amount_threshold = st.sidebar.number_input(
    "Round-dollar threshold",
    min_value=0,
    value=10000,
    step=1000,
)

after_hours_start = st.sidebar.number_input(
    "After-hours start hour",
    min_value=0,
    max_value=23,
    value=20,
)

after_hours_end = st.sidebar.number_input(
    "After-hours end hour",
    min_value=0,
    max_value=23,
    value=6,
)

dormant_days = st.sidebar.number_input(
    "Dormant account inactivity days",
    min_value=30,
    value=180,
    step=30,
)


# ============================================================
# Landing Page
# ============================================================

if oracle_file is None or coa_file is None:
    st.info("Upload both the Oracle Fusion export and the COA workbook to begin.")

    st.markdown(
        """
        ### Required Oracle Export Fields

        Your Oracle Fusion export should ideally include:

        - Account String / Code Combination / Concatenated Segments
        - Accounting Date / Transaction Date / GL Date
        - Description / Line Description
        - Debit and Credit, or Amount
        - Journal Source
        - Journal Category
        - Vendor / Supplier
        - Invoice Number
        - Batch ID / Batch Number / Batch Name
        - Inputter / Created By / Entered By / Prepared By
        - Approver / Approved By / Authorized By
        - Budget / Actual / Available Budget, if budget-overrun testing is needed
        """
    )

    st.stop()


# ============================================================
# Processing
# ============================================================

with st.spinner("Reading files and building dashboard..."):
    raw_df = read_excel_or_csv(oracle_file)
    lookups = load_coa(coa_file)

    df = raw_df.copy()

    account_string_col = detect_account_string_col(df)

    df = ensure_segment_columns(df)

    if account_string_col:
        df = split_account_string(df, account_string_col)

    df = detect_date_col(df)
    df = detect_amount_columns(df)
    df = add_coa_descriptions(df, lookups)

    df = add_exceptions(
        df,
        large_threshold=large_threshold,
        round_amount_threshold=round_amount_threshold,
        after_hours_start=after_hours_start,
        after_hours_end=after_hours_end,
        dormant_days=dormant_days,
    )

    df = add_duplicate_flags(df)

summary = create_summary(df)
exceptions = df[df["Has Exception"]].copy()
exception_summary = build_exception_summary(df)


# ============================================================
# Summaries
# ============================================================

dept_summary = (
    df.groupby(["Cost Center", "Cost Center Description"], dropna=False)["Absolute Amount"]
    .sum()
    .reset_index()
    .sort_values("Absolute Amount", ascending=False)
)

account_summary = (
    df.groupby(["Account", "Account Description"], dropna=False)["Absolute Amount"]
    .sum()
    .reset_index()
    .sort_values("Absolute Amount", ascending=False)
)

fund_summary = (
    df.groupby(["Fund", "Fund Description"], dropna=False)["Absolute Amount"]
    .sum()
    .reset_index()
    .sort_values("Absolute Amount", ascending=False)
)

program_summary = (
    df.groupby(["Program", "Program Description"], dropna=False)["Absolute Amount"]
    .sum()
    .reset_index()
    .sort_values("Absolute Amount", ascending=False)
)

top_dept = None
if not dept_summary.empty:
    r = dept_summary.iloc[0]
    top_dept = f"{r.get('Cost Center', '')} - {r.get('Cost Center Description', '')}"

top_account = None
if not account_summary.empty:
    r = account_summary.iloc[0]
    top_account = f"{r.get('Account', '')} - {r.get('Account Description', '')}"

conclusion = generate_conclusion(summary, top_dept, top_account)


# ============================================================
# Dashboard
# ============================================================

st.subheader("Executive KPI Summary")

k1, k2, k3, k4, k5 = st.columns(5)

k1.metric("Total Rows", f"{summary['Total Rows']:,}")
k2.metric("Gross Activity", f"${summary['Gross Activity']:,.0f}")
k3.metric("Net Amount", f"${summary['Net Amount']:,.0f}")
k4.metric("Exceptions", f"{summary['Exception Count']:,}")
k5.metric("Exception Rate", f"{summary['Exception Rate']:.1%}")

k6, k7, k8 = st.columns(3)

k6.metric("Total Debits", f"${summary['Total Debits']:,.0f}")
k7.metric("Total Credits", f"${summary['Total Credits']:,.0f}")
k8.metric("Potential Duplicates", f"{summary['Potential Duplicate Count']:,}")

st.subheader("Automated Analysis and Conclusion")
st.write(conclusion)


# ============================================================
# Exception Summary
# ============================================================

st.subheader("Exception Summary")

if exception_summary.empty:
    st.success("No exceptions were identified.")
else:
    c1, c2 = st.columns([1, 2])

    with c1:
        st.dataframe(exception_summary, use_container_width=True)

    with c2:
        fig = px.bar(
            exception_summary,
            x="Count",
            y="Exception Type",
            orientation="h",
            title="Exception Counts by Type",
        )
        fig.update_layout(yaxis={"categoryorder": "total ascending"})
        st.plotly_chart(fig, use_container_width=True)


# ============================================================
# Charts
# ============================================================

st.subheader("Interactive Charts")

tab1, tab2, tab3, tab4, tab5 = st.tabs(
    ["Cost Centers", "Accounts", "Funds", "Programs", "Monthly Trend"]
)

with tab1:
    top = dept_summary.head(15)

    fig = px.bar(
        top,
        x="Absolute Amount",
        y="Cost Center Description",
        orientation="h",
        title="Top Cost Centers by Gross Activity",
        hover_data=["Cost Center"],
    )

    fig.update_layout(yaxis={"categoryorder": "total ascending"})
    st.plotly_chart(fig, use_container_width=True)
    st.dataframe(dept_summary, use_container_width=True)

with tab2:
    top = account_summary.head(15)

    fig = px.bar(
        top,
        x="Absolute Amount",
        y="Account Description",
        orientation="h",
        title="Top Accounts by Gross Activity",
        hover_data=["Account"],
    )

    fig.update_layout(yaxis={"categoryorder": "total ascending"})
    st.plotly_chart(fig, use_container_width=True)
    st.dataframe(account_summary, use_container_width=True)

with tab3:
    top = fund_summary.head(15)

    fig = px.pie(
        top,
        names="Fund Description",
        values="Absolute Amount",
        title="Gross Activity by Fund",
    )

    st.plotly_chart(fig, use_container_width=True)
    st.dataframe(fund_summary, use_container_width=True)

with tab4:
    top = program_summary.head(15)

    fig = px.bar(
        top,
        x="Absolute Amount",
        y="Program Description",
        orientation="h",
        title="Top Programs by Gross Activity",
        hover_data=["Program"],
    )

    fig.update_layout(yaxis={"categoryorder": "total ascending"})
    st.plotly_chart(fig, use_container_width=True)
    st.dataframe(program_summary, use_container_width=True)

with tab5:
    if df["Transaction Date"].notna().any():
        monthly = df.dropna(subset=["Transaction Date"]).copy()
        monthly["Month"] = monthly["Transaction Date"].dt.to_period("M").astype(str)

        monthly = (
            monthly.groupby("Month")["Absolute Amount"]
            .sum()
            .reset_index()
            .sort_values("Month")
        )

        fig = px.line(
            monthly,
            x="Month",
            y="Absolute Amount",
            markers=True,
            title="Monthly Gross Activity Trend",
        )

        st.plotly_chart(fig, use_container_width=True)
        st.dataframe(monthly, use_container_width=True)
    else:
        st.warning("No valid transaction date was detected.")


# ============================================================
# Exceptions and Data
# ============================================================

st.subheader("Detailed Exception Report")

if exceptions.empty:
    st.success("No exceptions were identified using the current rules.")
else:
    st.warning(f"{len(exceptions):,} exception lines identified.")
    st.dataframe(exceptions, use_container_width=True)

st.subheader("Decoded Oracle Transactions")

display_df = exceptions if show_only_exceptions else df
st.dataframe(display_df, use_container_width=True)

with st.expander("COA Lookup Tables Used"):
    for name, lookup in lookups.items():
        st.markdown(f"### {name}")
        st.dataframe(lookup, use_container_width=True)


# ============================================================
# Exports
# ============================================================

st.subheader("Export Reports")

st.warning(
    "Before downloading reports, confirm that your device is secure. "
    "Exported Excel, Word, and PDF files may contain sensitive financial data."
)

excel_bytes = export_excel(df, summary, conclusion, exception_summary)
word_bytes = export_word(summary, conclusion, exceptions, exception_summary)
pdf_bytes = export_pdf(summary, conclusion, exceptions, exception_summary)

c1, c2, c3 = st.columns(3)

with c1:
    st.download_button(
        "Download Excel Report",
        data=excel_bytes,
        file_name="oracle_fusion_coa_audit_report.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )

with c2:
    st.download_button(
        "Download Word Report",
        data=word_bytes,
        file_name="oracle_fusion_coa_audit_report.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

with c3:
    st.download_button(
        "Download PDF Report",
        data=pdf_bytes,
        file_name="oracle_fusion_coa_audit_report.pdf",
        mime="application/pdf",
    )
